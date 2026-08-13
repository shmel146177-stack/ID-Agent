from pathlib import Path

from docx import Document

from app.generators.executive_doc_generator_v3 import (
    ExecutiveDocumentGeneratorV3,
)


def test_executive_doc_generator_v3_creates_real_docx(
    monkeypatch,
    tmp_path,
):

    monkeypatch.chdir(tmp_path)

    generator = ExecutiveDocumentGeneratorV3()

    data = {
        "equipment": 'Шкаф управления TEST:001/А',
        "manufacturer": "ООО Тест",
        "date": "12.08.2026",
        "drawing_number": "TEST-DWG-001",
        "power": "7,5 кВт",
        "current": "10-16 А",
        "voltage": "400 В",
        "ip": "IP66",
    }

    result = generator.create(data)

    output_path = Path(result)

    assert output_path.exists()
    assert output_path.is_file()
    assert output_path.suffix == ".docx"

    assert ":" not in output_path.name
    assert "/" not in output_path.name

    document = Document(output_path)

    assert len(document.tables) == 3

    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )

    assert "ID-Agent" in all_text
    assert 'Шкаф управления TEST:001/А' in all_text
    assert "ООО Тест" in all_text
    assert "TEST-DWG-001" in all_text
    assert "7,5 кВт" in all_text
    assert "IP66" in all_text

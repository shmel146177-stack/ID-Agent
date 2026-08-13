import json
from pathlib import Path

from docx import Document

import app.generators.hidden_works_journal_generator as generator_module
from app.generators.hidden_works_journal_generator import HiddenWorksJournalGenerator


def test_hidden_works_journal_uses_saved_act_data(
    monkeypatch,
    tmp_path,
):

    generator = HiddenWorksJournalGenerator()

    project_name = "TEST_PROJECT"
    project_path = tmp_path / project_name
    output_folder = project_path / "output"
    project_card_path = project_path / "project_card.json"
    act_data_path = project_path / "hidden_works_act_data.json"

    project_path.mkdir(parents=True)

    project_card_path.write_text(
        "{}",
        encoding="utf-8",
    )

    act_data_path.write_text(
        json.dumps(
            {
                "project": project_name,
                "acts": {
                    "grounding_device": {
                        "act_number": "А-005",
                        "act_date": "12.08.2026",
                        "work_start_date": "10.08.2026",
                        "work_finish_date": "11.08.2026",
                    }
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        generator,
        "_project_path",
        lambda name: project_path,
    )

    monkeypatch.setattr(
        generator,
        "_project_card_path",
        lambda name: project_card_path,
    )

    monkeypatch.setattr(
        generator,
        "_output_folder",
        lambda name: output_folder,
    )

    registry = {
        "project": project_name,
        "acts_count": 1,
        "high_priority_count": 1,
        "requires_field_confirmation": False,
        "acts": [
            {
                "code": "grounding_device",
                "act_title": "АОСР на устройство заземляющего устройства",
                "priority": "Высокий",
                "status": "Подтверждено",
                "evidence": [],
            }
        ],
    }

    monkeypatch.setattr(
        generator_module.hidden_works_registry,
        "analyze_project",
        lambda name: registry,
    )

    output_file = generator.create(project_name)

    document = Document(Path(output_file))

    text_parts = []

    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    full_text = "\n".join(text_parts)

    assert "А-005" in full_text
    assert "10.08.2026" in full_text
    assert "11.08.2026" in full_text

from pathlib import Path

from docx import Document

import app.generators.project_report_generator as report_module
from app.generators.project_report_generator import ProjectReportGenerator


def test_project_report_generator_creates_real_docx(monkeypatch, tmp_path):

    generator = ProjectReportGenerator()

    project_name = "TEST_PROJECT"
    output_path = tmp_path / "project_report.docx"

    project_card = {
        "object_name": "Тестовый объект",
        "address": "Тестовый адрес",
        "customer": "ООО Заказчик",
        "contractor": "ООО Подрядчик",
        "designer": "ООО Проектировщик",
        "contract_number": "TEST-001",
        "chief_engineer": "Иванов И.И.",
    }

    page_analysis = {
        "documents_count": 1,
        "pages_count": 2,
        "ocr_pages_count": 1,
        "documents": [
            {
                "filename": "project.pdf",
                "page_types": {
                    "Общие данные": 1,
                    "Электрическая схема": 1,
                },
            }
        ],
    }

    drawing_register = {
        "expected_sheet_count": 2,
        "registers": [
            {
                "entries": [
                    {
                        "sheet_number": 1,
                        "title": "Общие данные",
                        "number_source": "original",
                    },
                    {
                        "sheet_number": 2,
                        "title": "Электрическая схема",
                        "number_source": "restored_sequence",
                    },
                ]
            }
        ],
    }

    completeness = {
        "check_method": "Ведомость рабочих чертежей",
        "status": "Полный комплект",
        "required_count": 2,
        "found_count": 2,
        "missing_count": 0,
        "completeness_percent": 100.0,
        "documents": [
            {
                "sheet_number": 1,
                "title": "Общие данные",
                "status": "Есть",
                "matched_page": 1,
                "matched_page_type": "Общие данные",
                "confidence": "Высокая",
            },
            {
                "sheet_number": 2,
                "title": "Электрическая схема",
                "status": "Есть",
                "matched_page": 2,
                "matched_page_type": "Электрическая схема",
                "confidence": "Высокая",
            },
        ],
        "missing_sheets": [],
    }

    monkeypatch.setattr(
        generator,
        "_load_project_card",
        lambda name: project_card,
    )

    monkeypatch.setattr(
        generator,
        "_analysis_path",
        lambda name, filename: tmp_path / filename,
    )

    def fake_load_json(path, default=None):
        if Path(path).name == "page_analysis.json":
            return page_analysis

        if Path(path).name == "drawing_register.json":
            return drawing_register

        return default

    monkeypatch.setattr(
        generator,
        "_load_json",
        fake_load_json,
    )

    monkeypatch.setattr(
        report_module.document_completeness,
        "check",
        lambda name: completeness,
    )

    monkeypatch.setattr(
        generator,
        "_output_path",
        lambda name: output_path,
    )

    result = generator.create(project_name)

    assert result == str(output_path)

    assert output_path.exists()
    assert output_path.is_file()
    assert output_path.suffix.lower() == ".docx"
    assert output_path.stat().st_size > 0

    document = Document(output_path)

    assert len(document.paragraphs) > 0
    assert len(document.tables) >= 4

    full_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert "TEST_PROJECT" in full_text
    assert "ID-Agent" in full_text

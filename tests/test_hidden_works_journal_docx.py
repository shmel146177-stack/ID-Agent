from pathlib import Path

from docx import Document

import app.generators.hidden_works_journal_generator as generator_module
from app.generators.hidden_works_journal_generator import HiddenWorksJournalGenerator


def test_hidden_works_journal_generator_creates_real_docx(monkeypatch, tmp_path):

    generator = HiddenWorksJournalGenerator()

    project_name = "TEST_PROJECT"
    project_path = tmp_path / project_name
    output_folder = project_path / "output"
    project_card_path = project_path / "project_card.json"

    project_path.mkdir(parents=True)

    project_card_path.write_text(
        "{}",
        encoding="utf-8",
    )

    monkeypatch.setattr(
        generator,
        "_project_path",
        lambda name: project_path,
    )

    monkeypatch.setattr(
        generator,
        "_project_card_path",
        lambda name: project_card_path,
    )

    monkeypatch.setattr(
        generator,
        "_output_folder",
        lambda name: output_folder,
    )

    registry = {
        "project": project_name,
        "acts_count": 1,
        "high_priority_count": 1,
        "requires_field_confirmation": True,
        "acts": [
            {
                "code": "grounding_device",
                "title": "Устройство заземляющего устройства",
                "act_title": "АОСР на устройство заземляющего устройства",
                "status": "Требует подтверждения",
                "priority": "Высокий",
                "confidence": "Высокая",
                "reason": "Обнаружены проектные признаки заземления",
                "confirmation_required": True,
                "confirmation": "Подтвердить фактическое выполнение работ",
                "evidence": [
                    {
                        "page_type": "Заземление",
                        "pages_count": 1,
                        "source": "Постраничный анализ",
                    }
                ],
            }
        ],
    }

    monkeypatch.setattr(
        generator_module.hidden_works_registry,
        "analyze_project",
        lambda name: registry,
    )

    output_file = generator.create(
        project_name
    )

    output_path = Path(output_file)

    assert output_path.exists()
    assert output_path.is_file()
    assert output_path.suffix.lower() == ".docx"
    assert output_path.stat().st_size > 0

    document = Document(output_path)

    assert len(document.paragraphs) > 0
    assert len(document.tables) > 0

    full_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert "ID-Agent" in full_text

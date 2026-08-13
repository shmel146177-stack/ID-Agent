from pathlib import Path

from docx import Document

from app.generators.hidden_works_act_generator import HiddenWorksActGenerator


def test_hidden_works_act_generator_creates_real_docx(monkeypatch, tmp_path):

    generator = HiddenWorksActGenerator()

    project_name = "TEST_PROJECT"
    project_path = tmp_path / project_name
    output_folder = project_path / "output"
    project_card_path = project_path / "project_card.json"

    project_path.mkdir(parents=True)

    project_card_path.write_text(
        "{}",
        encoding="utf-8",
    )

    monkeypatch.setattr(
        generator,
        "_project_path",
        lambda name: project_path,
    )

    monkeypatch.setattr(
        generator,
        "_project_card_path",
        lambda name: project_card_path,
    )

    monkeypatch.setattr(
        generator,
        "_output_folder",
        lambda name: output_folder,
    )

    registry = {
        "acts_count": 1,
        "acts": [
            {
                "code": "grounding_device",
                "act_title": "АОСР на устройство заземляющего устройства",
                "priority": "Высокий",
                "evidence": [
                    {
                        "page_type": "Заземление",
                        "pages_count": 1,
                        "source": "Тест",
                    }
                ],
            }
        ],
    }

    output_file = generator.create(
        project_name,
        act_code="grounding_device",
        registry=registry,
        act_data={},
    )

    output_path = Path(output_file)

    assert output_path.exists()
    assert output_path.is_file()
    assert output_path.suffix.lower() == ".docx"
    assert output_path.stat().st_size > 0

    document = Document(output_path)

    assert len(document.paragraphs) > 0
    assert len(document.tables) > 0

    full_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert "ID-Agent" in full_text

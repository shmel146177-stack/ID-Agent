from pathlib import Path

from docx import Document

from app.generators.hidden_works_act_generator import HiddenWorksActGenerator


def test_hidden_works_act_docx_contains_extended_fields(
    monkeypatch,
    tmp_path,
):

    generator = HiddenWorksActGenerator()

    project_name = "TEST_PROJECT"
    project_path = tmp_path / project_name
    output_folder = project_path / "output"
    project_card_path = project_path / "project_card.json"

    project_path.mkdir(parents=True)

    project_card_path.write_text(
        "{}",
        encoding="utf-8",
    )

    monkeypatch.setattr(
        generator,
        "_project_path",
        lambda name: project_path,
    )

    monkeypatch.setattr(
        generator,
        "_project_card_path",
        lambda name: project_card_path,
    )

    monkeypatch.setattr(
        generator,
        "_output_folder",
        lambda name: output_folder,
    )

    registry = {
        "acts_count": 1,
        "acts": [
            {
                "code": "grounding_device",
                "act_title": (
                    "АОСР на устройство "
                    "заземляющего устройства"
                ),
                "priority": "Высокий",
                "evidence": [],
            }
        ],
    }

    act_data = {
        "customer_representative": "Customer representative",
        "contractor_representative": "Contractor representative",
        "construction_control_representative": "Control representative",
        "designer_representative": "Designer representative",
        "compliance": (
            "Работы соответствуют проектной документации"
        ),
        "next_works": "Обратная засыпка",
        "remarks": "Замечаний нет",
        "attachments": "Исполнительная схема ИС-01",
        "materials_compliance": "Materials comply with project",
        "test_results": "Test protocol 15",
        "geometric_parameters": "Elevation matches project",
        "actual_materials": (
            "Полоса стальная 40х5 мм, электроды L=3 м"
        ),
    }

    output_file = generator.create(
        project_name,
        act_code="grounding_device",
        registry=registry,
        act_data=act_data,
    )

    document = Document(Path(output_file))

    text_parts = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    full_text = "\n".join(text_parts)

    assert (
        "Работы соответствуют проектной документации"
        in full_text
    )
    assert "Обратная засыпка" in full_text
    assert "Замечаний нет" in full_text
    assert "Исполнительная схема ИС-01" in full_text
    assert (
        "Полоса стальная 40х5 мм, электроды L=3 м"
        in full_text
    )
    assert (
        "[ПОДТВЕРДИТЬ ФАКТИЧЕСКИ]"
        not in full_text
    )
    assert "Materials comply with project" in full_text
    assert "Test protocol 15" in full_text
    assert "Elevation matches project" in full_text
    assert (
        "[\u041f\u041e\u0414\u0422\u0412\u0415\u0420\u0414\u0418\u0422\u042c / "
        "\u0423\u0422\u041e\u0427\u041d\u0418\u0422\u042c]"
        not in full_text
    )
    assert (
        "[\u0423\u041a\u0410\u0417\u0410\u0422\u042c "
        "\u043f\u0440\u043e\u0442\u043e\u043a\u043e\u043b\u044b "
        "\u0438 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u044f]"
        not in full_text
    )
    assert (
        "[\u0423\u041a\u0410\u0417\u0410\u0422\u042c "
        "\u041f\u041e \u0424\u0410\u041a\u0422\u0423]"
        not in full_text
    )
    assert "[\u0424.\u0418.\u041e.]" not in full_text

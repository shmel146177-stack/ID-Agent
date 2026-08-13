import json
from pathlib import Path

from docx import Document

from app.generators.project_executive_generator import ProjectExecutiveGenerator


def test_project_executive_generator_creates_real_docx(monkeypatch, tmp_path):

    monkeypatch.chdir(tmp_path)

    project_name = "TEST_PROJECT"

    analysis_path = (
        tmp_path
        / "projects"
        / project_name
        / "analysis"
    )

    analysis_path.mkdir(
        parents=True,
        exist_ok=True,
    )

    project_analysis = {
        "documents": [
            {
                "filename": "drawing.pdf",
                "classification": "Чертеж",
                "status": "Тест",
                "analysis": {},
            }
        ]
    }

    completeness = {
        "status": "Тест",
        "completeness_percent": 50,
        "found_count": 1,
        "required_count": 2,
        "documents": [
            {
                "document_type": "Чертеж",
                "status": "Есть",
            },
            {
                "document_type": "Протокол",
                "status": "Нет",
            },
        ],
    }

    (
        analysis_path
        / "project_analysis.json"
    ).write_text(
        json.dumps(
            project_analysis,
            ensure_ascii=False,
            indent=4,
        ),
        encoding="utf-8",
    )

    (
        analysis_path
        / "document_completeness.json"
    ).write_text(
        json.dumps(
            completeness,
            ensure_ascii=False,
            indent=4,
        ),
        encoding="utf-8",
    )

    generator = ProjectExecutiveGenerator()

    result = generator.create(
        project_name
    )

    output_path = Path(result)

    assert output_path.exists()
    assert output_path.is_file()
    assert output_path.suffix == ".docx"

    document = Document(output_path)

    assert len(document.tables) == 4

    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )

    assert project_name in all_text
    assert "drawing.pdf" in all_text
    assert "50" in all_text

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.services.document_completeness import (
    document_completeness,
)


class ProjectReportGenerator:
    """
    Формирование итогового отчёта ID-Agent по проекту.

    В отчёт включаются:
    - карточка объекта;
    - сведения об обработанных документах;
    - постраничный анализ;
    - ведомость рабочих чертежей;
    - проверка комплектности;
    - отсутствующие листы;
    - итоговое заключение.
    """

    def _project_path(
        self,
        project_name: str,
    ) -> Path:

        return Path("projects") / project_name

    def _analysis_path(
        self,
        project_name: str,
        filename: str,
    ) -> Path:

        return self._project_path(project_name) / "analysis" / filename

    def _output_path(
        self,
        project_name: str,
    ) -> Path:

        return (
            self._project_path(project_name)
            / "output"
            / f"Отчет_ID-Agent_{project_name}.docx"
        )

    def _load_json(
        self,
        file_path: Path,
        default=None,
    ):

        if not file_path.exists():
            return default

        try:

            with open(
                file_path,
                "r",
                encoding="utf-8",
            ) as file:

                return json.load(file)

        except (
            OSError,
            json.JSONDecodeError,
        ):

            return default

    def _load_project_card(
        self,
        project_name: str,
    ) -> dict:

        project_file = self._project_path(project_name) / "project.json"

        return (
            self._load_json(
                project_file,
                {},
            )
            or {}
        )

    def _configure_document(
        self,
        document: Document,
    ):

        section = document.sections[0]

        section.top_margin = Cm(2)

        section.bottom_margin = Cm(2)

        section.left_margin = Cm(2)

        section.right_margin = Cm(1.5)

        styles = document.styles

        normal_style = styles["Normal"]

        normal_style.font.name = "Times New Roman"

        normal_style.font.size = Pt(11)

        for style_name in (
            "Title",
            "Heading 1",
            "Heading 2",
        ):

            style = styles[style_name]

            style.font.name = "Times New Roman"

    def _add_title(
        self,
        document: Document,
        project_name: str,
    ):

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("ИТОГОВЫЙ ОТЧЁТ ID-AGENT")

        run.bold = True

        run.font.name = "Times New Roman"

        run.font.size = Pt(16)

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(f"Проект: {project_name}")

        run.bold = True

        run.font.name = "Times New Roman"

        run.font.size = Pt(13)

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.add_run(
            "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M")
        )

    def _add_field_table(
        self,
        document: Document,
        rows: list[tuple[str, object]],
    ):

        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Параметр"

        header[1].text = "Значение"

        for cell in header:

            for run in cell.paragraphs[0].runs:

                run.bold = True

        for (
            label,
            value,
        ) in rows:

            cells = table.add_row().cells

            cells[0].text = str(label)

            cells[1].text = "" if value is None else str(value)

    def _add_project_information(
        self,
        document: Document,
        project_name: str,
        project_card: dict,
    ):

        document.add_heading(
            "1. Общие сведения об объекте",
            level=1,
        )

        rows = [
            (
                "Наименование проекта",
                project_name,
            ),
            (
                "Наименование объекта",
                project_card.get(
                    "object_name",
                    "",
                ),
            ),
            (
                "Адрес объекта",
                project_card.get(
                    "address",
                    "",
                ),
            ),
            (
                "Заказчик",
                project_card.get(
                    "customer",
                    "",
                ),
            ),
            (
                "Подрядчик",
                project_card.get(
                    "contractor",
                    "",
                ),
            ),
            (
                "Проектировщик",
                project_card.get(
                    "designer",
                    "",
                ),
            ),
            (
                "Номер договора",
                project_card.get(
                    "contract_number",
                    "",
                ),
            ),
            (
                "Главный инженер",
                project_card.get(
                    "chief_engineer",
                    "",
                ),
            ),
        ]

        self._add_field_table(
            document,
            rows,
        )

    def _add_processing_summary(
        self,
        document: Document,
        page_analysis: dict,
    ):

        document.add_heading(
            "2. Результаты анализа документов",
            level=1,
        )

        documents_count = page_analysis.get(
            "documents_count",
            0,
        )

        pages_count = page_analysis.get(
            "pages_count",
            0,
        )

        ocr_pages_count = page_analysis.get(
            "ocr_pages_count",
            0,
        )

        text_pages_count = pages_count - ocr_pages_count

        rows = [
            (
                "Обработано PDF-документов",
                documents_count,
            ),
            (
                "Всего страниц",
                pages_count,
            ),
            (
                "Страниц с текстовым слоем",
                text_pages_count,
            ),
            (
                "Страниц обработано OCR",
                ocr_pages_count,
            ),
        ]

        self._add_field_table(
            document,
            rows,
        )

        page_types = {}

        for document_data in page_analysis.get(
            "documents",
            [],
        ):

            for (
                page_type,
                count,
            ) in (
                document_data.get(
                    "page_types",
                    {},
                )
                or {}
            ).items():

                page_types[page_type] = (
                    page_types.get(
                        page_type,
                        0,
                    )
                    + count
                )

        if page_types:

            document.add_heading(
                "Типы обнаруженных страниц",
                level=2,
            )

            table = document.add_table(
                rows=1,
                cols=2,
            )

            table.style = "Table Grid"

            table.rows[0].cells[0].text = "Тип страницы"

            table.rows[0].cells[1].text = "Количество"

            for cell in table.rows[0].cells:

                for run in cell.paragraphs[0].runs:

                    run.bold = True

            for (
                page_type,
                count,
            ) in page_types.items():

                cells = table.add_row().cells

                cells[0].text = str(page_type)

                cells[1].text = str(count)

    def _add_drawing_register(
        self,
        document: Document,
        drawing_register: dict,
    ):

        document.add_heading(
            "3. Ведомость рабочих чертежей",
            level=1,
        )

        expected_sheet_count = drawing_register.get(
            "expected_sheet_count",
            0,
        )

        paragraph = document.add_paragraph()

        paragraph.add_run("Количество листов по ведомости: ").bold = True

        paragraph.add_run(str(expected_sheet_count))

        entries = []

        for register in drawing_register.get(
            "registers",
            [],
        ):

            entries.extend(
                register.get(
                    "entries",
                    [],
                )
            )

        if not entries:

            document.add_paragraph("Ведомость рабочих чертежей " "не обнаружена.")

            return

        table = document.add_table(
            rows=1,
            cols=3,
        )

        table.style = "Table Grid"

        headers = [
            "№ листа",
            "Наименование",
            "Источник номера",
        ]

        for index, header in enumerate(headers):

            cell = table.rows[0].cells[index]

            cell.text = header

            for run in cell.paragraphs[0].runs:

                run.bold = True

        for entry in entries:

            cells = table.add_row().cells

            cells[0].text = str(
                entry.get(
                    "sheet_number",
                    "",
                )
            )

            cells[1].text = str(
                entry.get(
                    "title",
                    "",
                )
            )

            number_source = entry.get("number_source")

            if number_source == "restored_sequence":

                source_text = "Восстановлен " "по последовательности"

            else:

                source_text = str(number_source or "")

            cells[2].text = source_text

    def _add_completeness(
        self,
        document: Document,
        completeness: dict,
    ):

        document.add_heading(
            "4. Проверка комплектности",
            level=1,
        )

        rows = [
            (
                "Метод проверки",
                completeness.get(
                    "check_method",
                    "",
                ),
            ),
            (
                "Статус",
                completeness.get(
                    "status",
                    "",
                ),
            ),
            (
                "Требуется листов",
                completeness.get(
                    "required_count",
                    0,
                ),
            ),
            (
                "Найдено листов",
                completeness.get(
                    "found_count",
                    0,
                ),
            ),
            (
                "Отсутствует листов",
                completeness.get(
                    "missing_count",
                    0,
                ),
            ),
            (
                "Комплектность",
                (f"{completeness.get('completeness_percent', 0)}%"),
            ),
        ]

        self._add_field_table(
            document,
            rows,
        )

        document.add_heading(
            "Сопоставление листов",
            level=2,
        )

        documents = completeness.get(
            "documents",
            [],
        )

        table = document.add_table(
            rows=1,
            cols=6,
        )

        table.style = "Table Grid"

        headers = [
            "№",
            "Наименование",
            "Статус",
            "PDF-стр.",
            "Тип страницы",
            "Уверенность",
        ]

        for index, header in enumerate(headers):

            cell = table.rows[0].cells[index]

            cell.text = header

            for run in cell.paragraphs[0].runs:

                run.bold = True

        for item in documents:

            cells = table.add_row().cells

            cells[0].text = str(
                item.get(
                    "sheet_number",
                    "",
                )
            )

            cells[1].text = str(
                item.get(
                    "title",
                    "",
                )
            )

            cells[2].text = str(
                item.get(
                    "status",
                    "",
                )
            )

            matched_page = item.get("matched_page")

            cells[3].text = "" if matched_page is None else str(matched_page)

            cells[4].text = str(
                item.get(
                    "matched_page_type",
                    "",
                )
                or ""
            )

            cells[5].text = str(
                item.get(
                    "confidence",
                    "",
                )
            )

    def _add_missing_sheets(
        self,
        document: Document,
        completeness: dict,
    ):

        document.add_heading(
            "5. Выявленные замечания",
            level=1,
        )

        missing_sheets = completeness.get(
            "missing_sheets",
            [],
        )

        if not missing_sheets:

            document.add_paragraph("Отсутствующие листы " "по ведомости не выявлены.")

            return

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Выявлены отсутствующие листы " "рабочей документации:")

        run.bold = True

        for sheet in missing_sheets:

            document.add_paragraph(
                (f"Лист №" f"{sheet.get('sheet_number')}: " f"{sheet.get('title')}"),
                style="List Bullet",
            )

    def _add_supporting_documents(
        self,
        document: Document,
        supporting_documents: dict,
    ):

        document.add_heading(
            "6. Сопроводительная исполнительная документация",
            level=1,
        )

        requirements_count = supporting_documents.get(
            "requirements_count",
            0,
        )

        high_priority_count = supporting_documents.get(
            "high_priority_count",
            0,
        )

        document.add_paragraph(
            f"Требуется документов: {requirements_count}. "
            f"Высокого приоритета: {high_priority_count}."
        )

        sections = supporting_documents.get(
            "sections",
            [],
        )

        if not sections:
            document.add_paragraph(
                "Сопроводительные документы автоматически не определены."
            )
            return

        for section in sections:

            number = section.get("number", "")
            title = section.get("title", "")
            required_count = section.get("required_count", 0)

            paragraph = document.add_paragraph()
            run = paragraph.add_run(
                f"Раздел {number}. {title} — требуется: {required_count}"
            )
            run.bold = True

            for item in section.get("documents", []):

                document_title = item.get(
                    "title",
                    item.get("code", ""),
                )

                priority = item.get(
                    "priority",
                    "",
                )

                text = document_title

                if priority:
                    text += f" (приоритет: {priority})"

                document.add_paragraph(
                    text,
                    style="List Bullet",
                )


    def _add_conclusion(
        self,
        document: Document,
        completeness: dict,
    ):

        document.add_heading(
            "6. Заключение ID-Agent",
            level=1,
        )

        required_count = completeness.get(
            "required_count",
            0,
        )

        found_count = completeness.get(
            "found_count",
            0,
        )

        missing_count = completeness.get(
            "missing_count",
            0,
        )

        percent = completeness.get(
            "completeness_percent",
            0,
        )

        if required_count > 0 and missing_count == 0:

            text = (
                "По результатам автоматического анализа "
                f"обнаружены все {required_count} листов, "
                "указанные в ведомости рабочих чертежей. "
                "Комплект проектной документации "
                "определён как полный. "
                f"Комплектность составляет {percent}%."
            )

        else:

            text = (
                "По результатам автоматического анализа "
                f"из {required_count} листов, указанных "
                "в ведомости рабочих чертежей, "
                f"обнаружено {found_count}. "
                f"Не обнаружено {missing_count}. "
                f"Расчётная комплектность составляет "
                f"{percent}%. "
                "Проект требует проверки отсутствующих "
                "листов перед использованием комплекта "
                "как полного."
            )

        document.add_paragraph(text)

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Примечание: ")

        run.bold = True

        paragraph.add_run(
            "отчёт сформирован автоматически "
            "на основании текста PDF, результатов OCR "
            "и сопоставления ведомости рабочих "
            "чертежей с найденными страницами. "
            "Результаты автоматического анализа "
            "рекомендуется подтверждать инженерной "
            "проверкой."
        )

    def create(
        self,
        project_name: str,
    ) -> str:

        project_card = self._load_project_card(project_name)

        page_analysis = (
            self._load_json(
                self._analysis_path(
                    project_name,
                    "page_analysis.json",
                ),
                {},
            )
            or {}
        )

        drawing_register = (
            self._load_json(
                self._analysis_path(
                    project_name,
                    "drawing_register.json",
                ),
                {},
            )
            or {}
        )

        supporting_documents = (
            self._load_json(
                self._analysis_path(
                    project_name,
                    "supporting_documents_registry.json",
                ),
                {},
            )
            or {}
        )

        completeness = document_completeness.check(project_name)

        document = Document()

        self._configure_document(document)

        self._add_title(
            document,
            project_name,
        )

        self._add_project_information(
            document,
            project_name,
            project_card,
        )

        self._add_processing_summary(
            document,
            page_analysis,
        )

        self._add_drawing_register(
            document,
            drawing_register,
        )

        self._add_completeness(
            document,
            completeness,
        )

        self._add_missing_sheets(
            document,
            completeness,
        )

        self._add_supporting_documents(
            document,
            supporting_documents,
        )

        self._add_conclusion(
            document,
            completeness,
        )

        output_path = self._output_path(project_name)

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(output_path)

        return str(output_path)


project_report_generator = ProjectReportGenerator()

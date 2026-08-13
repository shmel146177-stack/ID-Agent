from docx import Document
from docx.shared import Pt
import os


class ExecutiveDocumentGeneratorV2:

    def create(self, data: dict):

        os.makedirs("output", exist_ok=True)

        filename = "output/Исполнительная_документация_v2.docx"

        doc = Document()

        # Заголовок
        title = doc.add_heading(
            "ИСПОЛНИТЕЛЬНАЯ ДОКУМЕНТАЦИЯ",
            level=1
        )

        doc.add_paragraph(
            "Сформировано системой ID-Agent"
        )

        # Общие сведения
        doc.add_heading(
            "1. Общие сведения",
            level=2
        )

        table = doc.add_table(
            rows=4,
            cols=2
        )

        table.style = "Table Grid"

        table.cell(0,0).text = "Оборудование"
        table.cell(0,1).text = data.get("equipment", "")

        table.cell(1,0).text = "Изготовитель"
        table.cell(1,1).text = data.get("manufacturer", "")

        table.cell(2,0).text = "Дата"
        table.cell(2,1).text = data.get("date", "")

        table.cell(3,0).text = "Номер чертежа"
        table.cell(3,1).text = data.get("drawing_number", "")


        # Технические данные
        doc.add_heading(
            "2. Основные технические данные",
            level=2
        )

        doc.add_paragraph(
            data.get("equipment", "")
        )


        # Подписи
        doc.add_heading(
            "3. Подписи ответственных лиц",
            level=2
        )

        doc.add_paragraph(
            "Разработал: ____________________"
        )

        doc.add_paragraph(
            "Проверил: ______________________"
        )

        doc.add_paragraph(
            "Дата: __________________________"
        )


        doc.save(filename)

        return filename


executive_generator_v2 = ExecutiveDocumentGeneratorV2()
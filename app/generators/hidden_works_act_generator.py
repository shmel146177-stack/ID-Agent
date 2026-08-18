import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.services.hidden_works_registry import hidden_works_registry


class HiddenWorksActGenerator:
    """
    Генератор черновиков актов
    освидетельствования скрытых работ.

    Генератор НЕ подтверждает факт выполнения работ.

    Неизвестные фактические данные:
    - даты;
    - объёмы;
    - материалы;
    - участники;
    - протоколы;
    - исполнительные схемы

    автоматически не придумываются.
    """

    ACT_CONFIG = {
        "grounding_device": {
            "short_name": "заземление",
            "work_name": ("Устройство заземляющего устройства"),
            "next_works": (
                "Обратная засыпка, закрытие " "и дальнейшие монтажные работы"
            ),
            "materials_hint": (
                "Заземляющие электроды, полоса, " "проводники, соединительные элементы"
            ),
            "inspection_hint": (
                "Проверить расположение электродов, "
                "глубину, соединения, сварные швы, "
                "защитное покрытие и привязки."
            ),
            "attachments": [
                ("Исполнительная схема " "заземляющего устройства"),
                ("Протокол измерения сопротивления " "заземляющего устройства"),
                ("Паспорта / сертификаты " "применённых материалов"),
                "Фотофиксация скрытых работ",
            ],
        },
        "cable_entry": {
            "short_name": "кабельный_ввод",
            "work_name": ("Устройство скрытых участков " "кабельного ввода"),
            "next_works": (
                "Закрытие проходок, обратная засыпка, "
                "отделочные и последующие "
                "монтажные работы"
            ),
            "materials_hint": (
                "Кабель, защитные трубы, футляры, "
                "гильзы, проходки, элементы "
                "герметизации"
            ),
            "inspection_hint": (
                "Проверить трассу, глубину, "
                "защитные трубы и футляры, "
                "радиусы изгиба, проходки "
                "и герметизацию."
            ),
            "attachments": [
                ("Исполнительная схема " "кабельного ввода"),
                ("Документы качества на кабель " "и защитные трубы"),
                ("Протоколы испытаний кабельной " "линии — при необходимости"),
                "Фотофиксация скрытых работ",
            ],
        },
        "support_foundations": {
            "short_name": "временные_опоры",
            "work_name": ("Устройство оснований и скрытых " "элементов временных опор"),
            "next_works": ("Монтаж надземной части опор " "и последующие работы"),
            "materials_hint": (
                "Основания, закладные детали, "
                "крепёжные элементы, материалы "
                "обратной засыпки"
            ),
            "inspection_hint": (
                "Уточнить конструкцию опор. "
                "При наличии скрываемых оснований "
                "проверить размеры, глубину, "
                "положение и закрепление."
            ),
            "attachments": [
                ("Исполнительная схема " "расположения опор"),
                ("Документы качества " "на применённые материалы"),
                "Фотофиксация скрытых элементов",
            ],
        },
    }

    def _project_path(
        self,
        project_name: str,
    ) -> Path:

        return Path("projects") / project_name

    def _project_card_path(
        self,
        project_name: str,
    ) -> Path:

        return self._project_path(project_name) / "project.json"

    def _output_folder(
        self,
        project_name: str,
    ) -> Path:

        return (
            self._project_path(project_name)
            / "executive_docs"
            / "Исполнительная_документация"
            / "03_Акты_скрытых_работ"
        )

    def _load_json(
        self,
        path: Path,
    ) -> dict:

        if not path.exists():
            return {}

        with open(
            path,
            "r",
            encoding="utf-8",
        ) as file:

            return json.load(file)

    def _act_data_path(
        self,
        project_name: str,
    ) -> Path:

        return (
            self._project_path(project_name)
            / "hidden_works_act_data.json"
        )

    def load_act_data(
        self,
        project_name: str,
        act_code: str,
    ) -> dict:
        """
        ????????? ????? ?????????????? ???????????
        ?????? ??????????? ????.
        """

        data = self._load_json(
            self._act_data_path(project_name)
        )

        acts = data.get("acts", {})

        if not isinstance(acts, dict):
            return {}

        act_data = acts.get(act_code, {})

        if not isinstance(act_data, dict):
            return {}

        return dict(act_data)

    def save_act_data(
        self,
        project_name: str,
        act_code: str,
        act_data: dict,
    ) -> dict:
        """
        ????????? ?????? ?????????????? ???????????
        ???? ????.

        ?????? ???????? ?? ???????????? ? ?? ????????
        ??? ??????????? ??????.
        """

        allowed_fields = (
            "act_number",
            "act_date",
            "customer_representative",
            "contractor_representative",
            "construction_control_representative",
            "designer_representative",
            "work_location",
            "actual_materials",
            "work_start_date",
            "work_finish_date",
            "executive_scheme",
            "compliance",
            "materials_compliance",
            "test_results",
            "geometric_parameters",
            "next_works",
            "remarks",
            "attachments",
        )

        confirmed = {}

        for field in allowed_fields:

            value = act_data.get(field)

            if isinstance(value, str):
                value = value.strip()

            if value:
                confirmed[field] = value

        storage_path = self._act_data_path(
            project_name
        )

        storage = self._load_json(storage_path)

        if not isinstance(storage, dict):
            storage = {}

        acts = storage.get("acts", {})

        if not isinstance(acts, dict):
            acts = {}

        existing = acts.get(act_code, {})

        if not isinstance(existing, dict):
            existing = {}

        merged = dict(existing)
        merged.update(confirmed)

        acts[act_code] = merged

        storage["project"] = project_name
        storage["acts"] = acts

        storage_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        with open(
            storage_path,
            "w",
            encoding="utf-8",
        ) as file:

            json.dump(
                storage,
                file,
                ensure_ascii=False,
                indent=2,
            )

        return {
            "project": project_name,
            "act_code": act_code,
            "saved_fields": sorted(
                confirmed.keys()
            ),
            "file": str(storage_path),
        }

    def _get_value(
        self,
        data: dict,
        *keys: str,
    ) -> str | None:

        for key in keys:

            value = data.get(key)

            if value not in (
                None,
                "",
            ):
                return str(value)

        return None

    def _safe_filename(
        self,
        value: str,
    ) -> str:

        value = re.sub(
            r'[<>:"/\\|?*]',
            "_",
            value,
        )

        return value.strip(" .")

    def _configure_document(
        self,
        document: Document,
    ) -> None:

        section = document.sections[0]

        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)

        normal = document.styles["Normal"]

        normal.font.name = "Times New Roman"

        normal.font.size = Pt(10)

        normal._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

    def _set_cell_text(
        self,
        cell,
        text: str,
        bold: bool = False,
    ) -> None:

        cell.text = ""

        paragraph = cell.paragraphs[0]

        run = paragraph.add_run(str(text))

        run.bold = bold
        run.font.name = "Times New Roman"

        run.font.size = Pt(10)

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _shade_cell(
        self,
        cell,
        fill: str,
    ) -> None:

        tc_pr = cell._tc.get_or_add_tcPr()

        shading = OxmlElement("w:shd")

        shading.set(
            qn("w:fill"),
            fill,
        )

        tc_pr.append(shading)

    def _add_warning(
        self,
        document: Document,
    ) -> None:

        table = document.add_table(
            rows=1,
            cols=1,
        )

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(
            0,
            0,
        )

        self._shade_cell(
            cell,
            "FFF2CC",
        )

        self._set_cell_text(
            cell,
            (
                "ЧЕРНОВИК ID-AGENT. "
                "Не является подтверждением "
                "фактически выполненных работ. "
                "Перед подписанием необходимо "
                "проверить даты, объёмы, материалы, "
                "исполнительную документацию "
                "и полномочия участников."
            ),
            bold=True,
        )

        document.add_paragraph()

    def _add_title(
        self,
        document: Document,
    ) -> None:

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("АКТ\n" "ОСВИДЕТЕЛЬСТВОВАНИЯ " "СКРЫТЫХ РАБОТ")

        run.bold = True
        run.font.name = "Times New Roman"

        run.font.size = Pt(14)

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("ЧЕРНОВИК")

        run.bold = True
        run.font.size = Pt(12)

    def _add_act_number(
        self,
        document: Document,
        act_data: dict | None = None,
    ) -> None:

        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        act_data = act_data or {}

        act_number = (
            act_data.get("act_number")
            or "[УКАЗАТЬ]"
        )

        act_date = (
            act_data.get("act_date")
            or "[УКАЗАТЬ]"
        )

        self._set_cell_text(
            table.cell(0, 0),
            f"№ акта: {act_number}",
        )

        self._set_cell_text(
            table.cell(0, 1),
            f"Дата: {act_date}",
        )

        document.add_paragraph()

    def _add_object_information(
        self,
        document: Document,
        project_name: str,
        project_card: dict,
    ) -> None:

        object_name = (
            self._get_value(
                project_card,
                "object_name",
                "object",
                "name",
            )
            or project_name
        )

        address = (
            self._get_value(
                project_card,
                "address",
                "object_address",
            )
            or "[УКАЗАТЬ]"
        )

        customer = (
            self._get_value(
                project_card,
                "customer",
                "developer",
            )
            or "[УКАЗАТЬ]"
        )

        contractor = (
            self._get_value(
                project_card,
                "contractor",
                "general_contractor",
            )
            or "[УКАЗАТЬ]"
        )

        designer = (
            self._get_value(
                project_card,
                "designer",
                "project_organization",
            )
            or "[УКАЗАТЬ]"
        )

        values = [
            (
                "Объект",
                object_name,
            ),
            (
                "Адрес объекта",
                address,
            ),
            (
                "Заказчик / застройщик",
                customer,
            ),
            (
                "Подрядная организация",
                contractor,
            ),
            (
                "Проектная организация",
                designer,
            ),
        ]

        table = document.add_table(
            rows=len(values),
            cols=2,
        )

        table.style = "Table Grid"

        for row_index, (
            label,
            value,
        ) in enumerate(values):

            self._set_cell_text(
                table.cell(
                    row_index,
                    0,
                ),
                label,
                bold=True,
            )

            self._set_cell_text(
                table.cell(
                    row_index,
                    1,
                ),
                value,
            )

        document.add_paragraph()

    def _add_section_heading(
        self,
        document: Document,
        title: str,
    ) -> None:

        paragraph = document.add_paragraph()

        run = paragraph.add_run(title)

        run.bold = True
        run.font.name = "Times New Roman"

        run.font.size = Pt(11)

    def _add_participants(
        self,
        document: Document,
        act_data: dict | None = None,
    ) -> None:

        act_data = act_data or {}

        self._add_section_heading(
            document,
            (
                "\u041f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u0438\u0442\u0435\u043b\u0438, "
                "\u0443\u0447\u0430\u0441\u0442\u0432\u0443\u044e\u0449\u0438\u0435 "
                "\u0432 \u043e\u0441\u0432\u0438\u0434\u0435\u0442\u0435\u043b\u044c\u0441\u0442\u0432\u043e\u0432\u0430\u043d\u0438\u0438"
            ),
        )

        default_rep = (
            "[\u0424.\u0418.\u041e., \u0434\u043e\u043b\u0436\u043d\u043e\u0441\u0442\u044c, "
            "\u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442 \u043e \u043f\u043e\u043b\u043d\u043e\u043c\u043e\u0447\u0438\u044f\u0445]"
        )

        designer_default = (
            "[\u043f\u0440\u0438 \u043d\u0435\u043e\u0431\u0445\u043e\u0434\u0438\u043c\u043e\u0441\u0442\u0438: "
            "\u0424.\u0418.\u041e., \u0434\u043e\u043b\u0436\u043d\u043e\u0441\u0442\u044c]"
        )

        participants = [
            (
                "\u041f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u0438\u0442\u0435\u043b\u044c "
                "\u0437\u0430\u0441\u0442\u0440\u043e\u0439\u0449\u0438\u043a\u0430 / "
                "\u0437\u0430\u043a\u0430\u0437\u0447\u0438\u043a\u0430",
                act_data.get("customer_representative")
                or default_rep,
            ),
            (
                "\u041f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u0438\u0442\u0435\u043b\u044c "
                "\u0441\u0442\u0440\u043e\u0438\u0442\u0435\u043b\u044c\u043d\u043e\u0439 "
                "\u043e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u0438",
                act_data.get("contractor_representative")
                or default_rep,
            ),
            (
                "\u041f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u0438\u0442\u0435\u043b\u044c "
                "\u0441\u0442\u0440\u043e\u0438\u0442\u0435\u043b\u044c\u043d\u043e\u0433\u043e "
                "\u043a\u043e\u043d\u0442\u0440\u043e\u043b\u044f",
                act_data.get(
                    "construction_control_representative"
                )
                or default_rep,
            ),
            (
                "\u041f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u0438\u0442\u0435\u043b\u044c "
                "\u043f\u0440\u043e\u0435\u043a\u0442\u043d\u043e\u0439 "
                "\u043e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u0438",
                act_data.get("designer_representative")
                or designer_default,
            ),
        ]

        table = document.add_table(
            rows=len(participants),
            cols=2,
        )

        table.style = "Table Grid"

        for row_index, (
            role,
            value,
        ) in enumerate(participants):

            self._set_cell_text(
                table.cell(
                    row_index,
                    0,
                ),
                role,
                bold=True,
            )

            self._set_cell_text(
                table.cell(
                    row_index,
                    1,
                ),
                value,
            )

        document.add_paragraph()

    def _build_project_evidence(
        self,
        act: dict,
    ) -> list[str]:

        result = []

        for evidence in act.get(
            "evidence",
            [],
        ):

            sheet_number = evidence.get("sheet_number")

            title = evidence.get("title")

            page_type = evidence.get("page_type")

            pages_count = evidence.get("pages_count")

            if sheet_number is not None and title:

                result.append((f"Лист {sheet_number}: " f"{title}"))

            elif page_type:

                text = f"Тип страниц: {page_type}"

                if pages_count:

                    text += f" — {pages_count} стр."

                result.append(text)

        return result

    def _add_work_information(
        self,
        document: Document,
        act: dict,
        config: dict,
        act_data: dict | None = None,
    ) -> None:

        self._add_section_heading(
            document,
            "Освидетельствуемые работы",
        )

        evidence = self._build_project_evidence(act)

        evidence_text = (
            "\n".join(evidence)
            if evidence
            else "[УКАЗАТЬ]"
        )

        act_data = act_data or {}

        work_location = (
            act_data.get("work_location")
            or "[УКАЗАТЬ участок / оси / отметки / помещение]"
        )

        actual_materials_value = act_data.get(
            "actual_materials"
        )

        actual_materials = (
            actual_materials_value
            or "[УКАЗАТЬ марки, количество и документы качества]"
        )

        materials_hint = config["materials_hint"]

        if not actual_materials_value:
            materials_hint += (
                ". [ПОДТВЕРДИТЬ ФАКТИЧЕСКИ]"
            )

        work_start_date = (
            act_data.get("work_start_date")
            or "[УКАЗАТЬ]"
        )

        work_finish_date = (
            act_data.get("work_finish_date")
            or "[УКАЗАТЬ]"
        )

        executive_scheme = (
            act_data.get("executive_scheme")
            or "[УКАЗАТЬ номер и наименование либо приложить]"
        )

        values = [
            (
                "Наименование работ",
                config["work_name"],
            ),
            (
                "Место выполнения работ",
                work_location,
            ),
            (
                "Проектная документация",
                evidence_text,
            ),
            (
                "Предполагаемые материалы",
                materials_hint,
            ),
            (
                "Фактически применённые материалы и изделия",
                actual_materials,
            ),
            (
                "Дата начала работ",
                work_start_date,
            ),
            (
                "Дата окончания работ",
                work_finish_date,
            ),
            (
                "Исполнительная схема",
                executive_scheme,
            ),
        ]

        table = document.add_table(
            rows=len(values),
            cols=2,
        )

        table.style = "Table Grid"

        for row_index, (
            label,
            value,
        ) in enumerate(values):

            self._set_cell_text(
                table.cell(
                    row_index,
                    0,
                ),
                label,
                bold=True,
            )

            self._set_cell_text(
                table.cell(
                    row_index,
                    1,
                ),
                value,
            )

        document.add_paragraph()

        self._add_section_heading(
            document,
            "Что необходимо проверить",
        )

        document.add_paragraph(
            config["inspection_hint"]
        )

    def _add_inspection_result(
        self,
        document: Document,
        config: dict,
        act_data: dict | None = None,
    ) -> None:

        act_data = act_data or {}

        compliance = (
            act_data.get("compliance")
            or "[\u041f\u041e\u0414\u0422\u0412\u0415\u0420\u0414\u0418\u0422\u042c / \u0423\u0422\u041e\u0427\u041d\u0418\u0422\u042c]"
        )

        materials_compliance = (
            act_data.get("materials_compliance")
            or (
                "[\u041f\u041e\u0414\u0422\u0412\u0415\u0420\u0414\u0418\u0422\u042c / "
                "\u0423\u0422\u041e\u0427\u041d\u0418\u0422\u042c]"
            )
        )

        test_results = (
            act_data.get("test_results")
            or (
                "[\u0423\u041a\u0410\u0417\u0410\u0422\u042c "
                "\u043f\u0440\u043e\u0442\u043e\u043a\u043e\u043b\u044b "
                "\u0438 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u044f]"
            )
        )

        geometric_parameters = (
            act_data.get("geometric_parameters")
            or (
                "[\u0423\u041a\u0410\u0417\u0410\u0422\u042c "
                "\u041f\u041e \u0424\u0410\u041a\u0422\u0423]"
            )
        )

        remarks = (
            act_data.get("remarks")
            or "[\u041d\u0415\u0422 / \u0423\u041a\u0410\u0417\u0410\u0422\u042c]"
        )

        self._add_section_heading(
            document,
            "\u0420\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u044b \u043e\u0441\u0432\u0438\u0434\u0435\u0442\u0435\u043b\u044c\u0441\u0442\u0432\u043e\u0432\u0430\u043d\u0438\u044f",
        )

        paragraphs = [
            (
                "1. \u0421\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0438\u0435 "
                "\u0432\u044b\u043f\u043e\u043b\u043d\u0435\u043d\u043d\u044b\u0445 \u0440\u0430\u0431\u043e\u0442 "
                "\u043f\u0440\u043e\u0435\u043a\u0442\u043d\u043e\u0439 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u0438: "
                f"{compliance}."
            ),
            (
                "2. \u0421\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0438\u0435 \u0444\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0438 "
                "\u043f\u0440\u0438\u043c\u0435\u043d\u0451\u043d\u043d\u044b\u0445 \u043c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u043e\u0432 \u043f\u0440\u043e\u0435\u043a\u0442\u0443 "
                "\u0438 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u043c \u043a\u0430\u0447\u0435\u0441\u0442\u0432\u0430: "
                f"{materials_compliance}."
            ),
            (
                "3. \u0420\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u044b \u0438\u0437\u043c\u0435\u0440\u0435\u043d\u0438\u0439 "
                "\u0438 \u0438\u0441\u043f\u044b\u0442\u0430\u043d\u0438\u0439: "
                f"{test_results}."
            ),
            (
                "4. \u0413\u0435\u043e\u043c\u0435\u0442\u0440\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u043f\u0430\u0440\u0430\u043c\u0435\u0442\u0440\u044b, "
                "\u0433\u043b\u0443\u0431\u0438\u043d\u044b, \u043e\u0442\u043c\u0435\u0442\u043a\u0438 \u0438 \u043f\u0440\u0438\u0432\u044f\u0437\u043a\u0438: "
                f"{geometric_parameters}."
            ),
            (
                "5. \u0417\u0430\u043c\u0435\u0447\u0430\u043d\u0438\u044f \u043f\u0440\u0438 "
                "\u043e\u0441\u0432\u0438\u0434\u0435\u0442\u0435\u043b\u044c\u0441\u0442\u0432\u043e\u0432\u0430\u043d\u0438\u0438: "
                f"{remarks}."
            ),
        ]

        for paragraph_text in paragraphs:
            document.add_paragraph(
                paragraph_text
            )

        document.add_paragraph()

        self._add_section_heading(
            document,
            "\u0420\u0430\u0437\u0440\u0435\u0448\u0435\u043d\u0438\u0435 \u043d\u0430 \u043f\u043e\u0441\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0435 \u0440\u0430\u0431\u043e\u0442\u044b",
        )

        next_works = act_data.get(
            "next_works"
        )

        if next_works:
            next_works_text = next_works
            confirmation_text = ""
        else:
            next_works_text = config[
                "next_works"
            ]
            confirmation_text = (
                " [\u041f\u041e\u0414\u0422\u0412\u0415\u0420\u0414\u0418\u0422\u042c "
                "\u041f\u0415\u0420\u0415\u0414 \u041f\u041e\u0414\u041f\u0418\u0421\u0410\u041d\u0418\u0415\u041c]"
            )

        document.add_paragraph(
            (
                "\u041d\u0430 \u043e\u0441\u043d\u043e\u0432\u0430\u043d\u0438\u0438 \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u043e\u0432 "
                "\u043e\u0441\u0432\u0438\u0434\u0435\u0442\u0435\u043b\u044c\u0441\u0442\u0432\u043e\u0432\u0430\u043d\u0438\u044f \u0440\u0430\u0437\u0440\u0435\u0448\u0430\u0435\u0442\u0441\u044f "
                "\u043f\u0440\u043e\u0438\u0437\u0432\u043e\u0434\u0441\u0442\u0432\u043e \u043f\u043e\u0441\u043b\u0435\u0434\u0443\u044e\u0449\u0438\u0445 \u0440\u0430\u0431\u043e\u0442: "
                f"{next_works_text}."
                f"{confirmation_text}"
            )
        )

    def _add_attachments(
        self,
        document: Document,
        config: dict,
        act_data: dict | None = None,
    ) -> None:

        act_data = act_data or {}

        self._add_section_heading(
            document,
            (
                "\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u044f \u0438 "
                "\u043f\u043e\u0434\u0442\u0432\u0435\u0440\u0436\u0434\u0430\u044e\u0449\u0438\u0435 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u044b"
            ),
        )

        attachments = act_data.get(
            "attachments"
        )

        if attachments:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.add_run(
                attachments
            )

        else:

            for item in config.get(
                "attachments",
                [],
            ):

                paragraph = document.add_paragraph(
                    style="List Bullet"
                )

                paragraph.add_run(
                    f"{item}: [\u0423\u041a\u0410\u0417\u0410\u0422\u042c]"
                )

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.add_run(
                "\u0418\u043d\u044b\u0435 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u044b: "
                "[\u0423\u041a\u0410\u0417\u0410\u0422\u042c]"
            )

    def _add_signatures(
        self,
        document: Document,
        act_data: dict | None = None,
    ) -> None:

        act_data = act_data or {}

        document.add_paragraph()

        self._add_section_heading(
            document,
            "Подписи участников",
        )

        rows = [
            ("Представитель " "заказчика / застройщика"),
            ("Представитель " "строительной организации"),
            ("Представитель " "строительного контроля"),
            ("Представитель " "проектной организации"),
        ]

        fields = [
            "customer_representative",
            "contractor_representative",
            "construction_control_representative",
            "designer_representative",
        ]

        table = document.add_table(
            rows=len(rows),
            cols=3,
        )

        table.style = "Table Grid"

        for index, (role, field) in enumerate(zip(rows, fields)):

            self._set_cell_text(
                table.cell(
                    index,
                    0,
                ),
                role,
            )

            self._set_cell_text(
                table.cell(
                    index,
                    1,
                ),
                "________________",
            )

            self._set_cell_text(
                table.cell(
                    index,
                    2,
                ),
                act_data.get(field) or "[\u0424.\u0418.\u041e.]",
            )

    def _find_act(
        self,
        registry: dict,
        act_code: str,
    ) -> dict | None:

        for act in registry.get(
            "acts",
            [],
        ):

            if act.get("code") == act_code:

                return act

        return None

    def create(
        self,
        project_name: str,
        act_code: str = "grounding_device",
        registry: dict | None = None,
        act_data: dict | None = None,
    ) -> str:

        project_path = self._project_path(project_name)

        if not project_path.exists():

            raise FileNotFoundError("Проект не найден: " f"{project_name}")

        if act_code not in self.ACT_CONFIG:

            raise ValueError("Неизвестный тип АОСР: " f"{act_code}")

        if registry is None:

            registry = hidden_works_registry.analyze_project(project_name)

        act = self._find_act(
            registry,
            act_code,
        )

        if act is None:

            raise ValueError(
                ("АОСР данного типа " "не определён для проекта: " f"{act_code}")
            )

        config = self.ACT_CONFIG[act_code]

        project_card = self._load_json(self._project_card_path(project_name))

        document = Document()

        self._configure_document(document)

        self._add_warning(document)

        self._add_title(document)

        self._add_act_number(
            document,
            act_data=act_data,
        )

        self._add_object_information(
            document,
            project_name,
            project_card,
        )

        self._add_participants(
            document,
            act_data=act_data,
        )

        self._add_work_information(
            document,
            act,
            config,
            act_data=act_data,
        )

        self._add_inspection_result(
            document,
            config,
            act_data=act_data,
        )

        self._add_attachments(
            document,
            config,
            act_data=act_data,
        )

        self._add_signatures(
            document,
            act_data=act_data,
        )

        document.add_paragraph()

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run(
            (
                "Черновик сформирован ID-Agent: "
                + datetime.now().strftime("%d.%m.%Y %H:%M")
            )
        )

        run.italic = True
        run.font.size = Pt(8)

        output_folder = self._output_folder(project_name)

        output_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        filename = (
            "АОСР_"
            + config["short_name"]
            + "_"
            + self._safe_filename(project_name)
            + ".docx"
        )

        output_file = output_folder / filename

        document.save(output_file)

        return str(output_file)

    def create_all(
        self,
        project_name: str,
    ) -> dict:
        """
        Создаёт все черновики АОСР,
        которые определил hidden_works_registry.
        """

        registry = hidden_works_registry.analyze_project(project_name)

        created = []
        skipped = []

        for act in registry.get(
            "acts",
            [],
        ):

            act_code = act.get("code")

            if act_code not in self.ACT_CONFIG:

                skipped.append(
                    {
                        "code": act_code,
                        "reason": ("Для данного типа " "пока нет шаблона DOCX"),
                    }
                )

                continue

            saved_act_data = self.load_act_data(
                project_name,
                act_code,
            )

            output_file = self.create(
                project_name,
                act_code,
                registry=registry,
                act_data=(
                    saved_act_data
                    if saved_act_data
                    else None
                ),
            )

            created.append(
                {
                    "code": act_code,
                    "title": (act.get("act_title")),
                    "priority": (act.get("priority")),
                    "file": output_file,
                }
            )

        return {
            "project": project_name,
            "status": "Готово",
            "acts_detected": (
                registry.get(
                    "acts_count",
                    0,
                )
            ),
            "acts_created": len(created),
            "acts_skipped": len(skipped),
            "created": created,
            "skipped": skipped,
            "requires_field_confirmation": (
                registry.get(
                    "requires_field_confirmation",
                    False,
                )
            ),
        }


hidden_works_act_generator = HiddenWorksActGenerator()

from docx import Document
import os


class ExecutiveDocumentGenerator:

    def create(self, data: dict):

        os.makedirs("output", exist_ok=True)

        filename = "output/Исполнительная_документация.docx"

        doc = Document()

        doc.add_heading(
            "Исполнительная документация",
            level=1
        )

        doc.add_paragraph(
            "Документ сформирован системой ID-Agent"
        )

        doc.add_heading(
            "Сведения об оборудовании",
            level=2
        )

        doc.add_paragraph(
            f"Оборудование: {data.get('equipment', '')}"
        )

        doc.add_paragraph(
            f"Изготовитель: {data.get('manufacturer', '')}"
        )

        doc.add_paragraph(
            f"Дата изготовления: {data.get('date', '')}"
        )

        doc.add_paragraph(
            f"Номер чертежа: {data.get('drawing_number', '')}"
        )

        doc.add_heading(
            "Подписи",
            level=2
        )

        doc.add_paragraph(
            "Ответственный представитель ____________"
        )

        doc.add_paragraph(
            "Дата ____________"
        )

        doc.save(filename)

        return filename


executive_generator = ExecutiveDocumentGenerator()
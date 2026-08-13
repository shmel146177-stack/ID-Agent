import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.services.hidden_works_registry import hidden_works_registry


class HiddenWorksJournalGenerator:
    """
    Генератор черновика журнала учёта АОСР.

    Журнал формируется на основании проектной
    документации и результатов анализа ID-Agent.

    Генератор не подтверждает:
    - факт выполнения работ;
    - даты;
    - номера актов;
    - объёмы;
    - участников;
    - подписание АОСР.

    Эти сведения должны быть заполнены
    и подтверждены по факту строительства.
    """

    def _project_path(
        self,
        project_name: str,
    ) -> Path:

        return Path("projects") / project_name

    def _project_card_path(
        self,
        project_name: str,
    ) -> Path:

        return self._project_path(project_name) / "project.json"

    def _output_folder(
        self,
        project_name: str,
    ) -> Path:

        return (
            self._project_path(project_name)
            / "executive_docs"
            / "Исполнительная_документация"
            / "07_Журналы_работ"
        )

    def _load_json(
        self,
        path: Path,
    ) -> dict:

        if not path.exists():
            return {}

        with open(
            path,
            "r",
            encoding="utf-8",
        ) as file:

            return json.load(file)

    def _get_value(
        self,
        data: dict,
        *keys: str,
    ) -> str | None:

        for key in keys:

            value = data.get(key)

            if value not in (
                None,
                "",
            ):

                return str(value)

        return None

    def _safe_filename(
        self,
        value: str,
    ) -> str:

        value = re.sub(
            r'[<>:"/\\|?*]',
            "_",
            value,
        )

        return value.strip(" .")

    def _configure_document(
        self,
        document: Document,
    ) -> None:

        section = document.sections[0]

        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        normal = document.styles["Normal"]

        normal.font.name = "Times New Roman"

        normal.font.size = Pt(9)

        normal._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

    def _set_cell_text(
        self,
        cell,
        text: str,
        bold: bool = False,
        size: int = 9,
    ) -> None:

        cell.text = ""

        paragraph = cell.paragraphs[0]

        run = paragraph.add_run(str(text))

        run.bold = bold

        run.font.name = "Times New Roman"

        run.font.size = Pt(size)

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _shade_cell(
        self,
        cell,
        fill: str,
    ) -> None:

        tc_pr = cell._tc.get_or_add_tcPr()

        shading = OxmlElement("w:shd")

        shading.set(
            qn("w:fill"),
            fill,
        )

        tc_pr.append(shading)

    def _add_warning(
        self,
        document: Document,
    ) -> None:

        table = document.add_table(
            rows=1,
            cols=1,
        )

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(
            0,
            0,
        )

        self._shade_cell(
            cell,
            "FFF2CC",
        )

        self._set_cell_text(
            cell,
            (
                "ЧЕРНОВИК ID-AGENT. "
                "Журнал сформирован автоматически "
                "по проектной документации. "
                "Наличие записи не подтверждает "
                "фактическое выполнение работ. "
                "Перед использованием необходимо "
                "заполнить фактические даты, номера "
                "актов и проверить документацию."
            ),
            bold=True,
            size=9,
        )

        document.add_paragraph()

    def _add_title(
        self,
        document: Document,
    ) -> None:

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(
            "ЖУРНАЛ УЧЁТА\n" "АКТОВ ОСВИДЕТЕЛЬСТВОВАНИЯ " "СКРЫТЫХ РАБОТ"
        )

        run.bold = True

        run.font.name = "Times New Roman"

        run.font.size = Pt(14)

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Times New Roman",
        )

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("ЧЕРНОВИК")

        run.bold = True
        run.font.size = Pt(11)

    def _add_project_information(
        self,
        document: Document,
        project_name: str,
        project_card: dict,
    ) -> None:

        object_name = (
            self._get_value(
                project_card,
                "object_name",
                "object",
                "name",
            )
            or project_name
        )

        address = (
            self._get_value(
                project_card,
                "address",
                "object_address",
            )
            or "[УКАЗАТЬ]"
        )

        customer = (
            self._get_value(
                project_card,
                "customer",
                "developer",
            )
            or "[УКАЗАТЬ]"
        )

        contractor = (
            self._get_value(
                project_card,
                "contractor",
                "general_contractor",
            )
            or "[УКАЗАТЬ]"
        )

        values = [
            (
                "Проект",
                project_name,
            ),
            (
                "Объект",
                object_name,
            ),
            (
                "Адрес",
                address,
            ),
            (
                "Заказчик / застройщик",
                customer,
            ),
            (
                "Подрядная организация",
                contractor,
            ),
        ]

        table = document.add_table(
            rows=len(values),
            cols=2,
        )

        table.style = "Table Grid"

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_index, (
            label,
            value,
        ) in enumerate(values):

            self._set_cell_text(
                table.cell(
                    row_index,
                    0,
                ),
                label,
                bold=True,
            )

            self._set_cell_text(
                table.cell(
                    row_index,
                    1,
                ),
                value,
            )

        document.add_paragraph()

    def _evidence_text(
        self,
        act: dict,
    ) -> str:

        result = []

        for evidence in act.get(
            "evidence",
            [],
        ):

            sheet_number = evidence.get("sheet_number")

            title = evidence.get("title")

            page_type = evidence.get("page_type")

            pages_count = evidence.get("pages_count")

            if sheet_number is not None and title:

                result.append((f"Лист {sheet_number}: " f"{title}"))

                continue

            if page_type:

                text = f"{page_type}"

                if pages_count:

                    text += f" ({pages_count} стр.)"

                result.append(text)

        if not result:

            return "Основание автоматически " "не определено"

        return "\n".join(result)

    def _add_registry_table(
        self,
        document: Document,
        registry: dict,
    ) -> None:

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Реестр актов")

        run.bold = True
        run.font.size = Pt(11)

        acts = registry.get(
            "acts",
            [],
        )

        headers = [
            "№",
            "Наименование АОСР",
            "Проектное основание",
            "Приоритет",
            "Дата работ",
            "№ акта",
            "Статус",
        ]

        table = document.add_table(
            rows=1 + len(acts),
            cols=len(headers),
        )

        table.style = "Table Grid"

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for column_index, header in enumerate(headers):

            cell = table.cell(
                0,
                column_index,
            )

            self._shade_cell(
                cell,
                "D9EAF7",
            )

            self._set_cell_text(
                cell,
                header,
                bold=True,
                size=8,
            )

        for row_index, act in enumerate(
            acts,
            start=1,
        ):

            values = [
                str(row_index),
                (act.get("act_title") or act.get("title") or ""),
                self._evidence_text(act),
                (act.get("priority") or ""),
                "[УКАЗАТЬ]",
                "[УКАЗАТЬ]",
                (act.get("status") or "Требует подтверждения"),
            ]

            for column_index, value in enumerate(values):

                self._set_cell_text(
                    table.cell(
                        row_index,
                        column_index,
                    ),
                    value,
                    size=8,
                )

        document.add_paragraph()

    def _add_summary(
        self,
        document: Document,
        registry: dict,
    ) -> None:

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Сводная информация")

        run.bold = True
        run.font.size = Pt(11)

        acts_count = registry.get(
            "acts_count",
            0,
        )

        high_priority_count = registry.get(
            "high_priority_count",
            0,
        )

        values = [
            (
                "АОСР определено",
                acts_count,
            ),
            (
                "Высокий приоритет",
                high_priority_count,
            ),
            (
                "Требует подтверждения по факту",
                (
                    "Да"
                    if registry.get(
                        "requires_field_confirmation",
                        False,
                    )
                    else "Нет"
                ),
            ),
        ]

        table = document.add_table(
            rows=len(values),
            cols=2,
        )

        table.style = "Table Grid"

        for row_index, (
            label,
            value,
        ) in enumerate(values):

            self._set_cell_text(
                table.cell(
                    row_index,
                    0,
                ),
                label,
                bold=True,
            )

            self._set_cell_text(
                table.cell(
                    row_index,
                    1,
                ),
                value,
            )

        document.add_paragraph()

    def _add_notes(
        self,
        document: Document,
    ) -> None:

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Порядок заполнения")

        run.bold = True
        run.font.size = Pt(11)

        notes = [
            (
                "Дата работ заполняется по фактическому "
                "выполнению скрываемой операции."
            ),
            ("Номер АОСР указывается после " "оформления соответствующего акта."),
            (
                "Проектное основание необходимо "
                "сверить с действующей рабочей "
                "документацией и изменениями проекта."
            ),
            (
                "Записи ID-Agent являются "
                "предварительными и требуют "
                "инженерной проверки."
            ),
        ]

        for text in notes:

            document.add_paragraph(
                text,
                style="List Bullet",
            )

    def create(
        self,
        project_name: str,
    ) -> str:

        project_path = self._project_path(project_name)

        if not project_path.exists():

            raise FileNotFoundError("Проект не найден: " f"{project_name}")

        # ---------------------------------------------------------
        # 1. РЕЕСТР АОСР
        # ---------------------------------------------------------

        registry = hidden_works_registry.analyze_project(project_name)

        # ---------------------------------------------------------
        # 2. КАРТОЧКА ПРОЕКТА
        # ---------------------------------------------------------

        project_card = self._load_json(self._project_card_path(project_name))

        # ---------------------------------------------------------
        # 3. DOCX
        # ---------------------------------------------------------

        document = Document()

        self._configure_document(document)

        self._add_warning(document)

        self._add_title(document)

        self._add_project_information(
            document,
            project_name,
            project_card,
        )

        self._add_summary(
            document,
            registry,
        )

        self._add_registry_table(
            document,
            registry,
        )

        self._add_notes(document)

        document.add_paragraph()

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run(
            (
                "Черновик сформирован ID-Agent: "
                + datetime.now().strftime("%d.%m.%Y %H:%M")
            )
        )

        run.italic = True
        run.font.size = Pt(8)

        # ---------------------------------------------------------
        # 4. СОХРАНЕНИЕ
        # ---------------------------------------------------------

        output_folder = self._output_folder(project_name)

        output_folder.mkdir(
            parents=True,
            exist_ok=True,
        )

        filename = "Журнал_скрытых_работ_" + self._safe_filename(project_name) + ".docx"

        output_file = output_folder / filename

        document.save(output_file)

        return str(output_file)


hidden_works_journal_generator = HiddenWorksJournalGenerator()

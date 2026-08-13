from docx import Document
from docx.shared import Pt
import os
import re


class ExecutiveDocumentGeneratorV4:

    def create(self, data: dict):

        os.makedirs("output", exist_ok=True)

        equipment = data.get(
            "equipment",
            "Оборудование"
        )

        safe_name = re.sub(
            r'[\\/*?:"<>|]',
            "",
            equipment
        )

        filename = f"output/ИД_{safe_name[:50]}.docx"


        doc = Document()


        # Настройка шрифта
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)



        # -------------------------
        # Титульный лист
        # -------------------------

        doc.add_heading(
            "ИСПОЛНИТЕЛЬНАЯ ДОКУМЕНТАЦИЯ",
            level=1
        )

        doc.add_paragraph(
            "Сформировано системой ID-Agent"
        )

        doc.add_paragraph(
            ""
        )

        doc.add_paragraph(
            "Объект: ______________________________"
        )

        doc.add_paragraph(
            f"Оборудование: {equipment}"
        )

        doc.add_paragraph(
            f"Дата изготовления: {data.get('date','')}"
        )



        # -------------------------
        # 1 Ведомость
        # -------------------------

        doc.add_heading(
            "1. Ведомость исполнительной документации",
            level=2
        )


        table = doc.add_table(
            rows=4,
            cols=3
        )

        table.style = "Table Grid"


        headers = [
            "№",
            "Наименование документа",
            "Обозначение"
        ]


        for i, h in enumerate(headers):
            table.cell(0, i).text = h


        rows = [

            (
                "1",
                "Паспорт оборудования",
                data.get("drawing_number", "")
            ),

            (
                "2",
                "Заводская документация",
                "PDF"
            ),

            (
                "3",
                "Исполнительная документация",
                "ID-Agent"
            )

        ]


        for r, row in enumerate(rows, start=1):

            for c, value in enumerate(row):

                table.cell(r, c).text = value



        # -------------------------
        # 2 Общие сведения
        # -------------------------

        doc.add_heading(
            "2. Общие сведения об оборудовании",
            level=2
        )


        table = doc.add_table(
            rows=5,
            cols=2
        )

        table.style = "Table Grid"


        info = [

            (
                "Оборудование",
                data.get("equipment","")
            ),

            (
                "Изготовитель",
                data.get("manufacturer","")
            ),

            (
                "Дата изготовления",
                data.get("date","")
            ),

            (
                "Номер чертежа",
                data.get("drawing_number","")
            ),

            (
                "Тип документа",
                data.get("document_type","")
            )

        ]


        for i,row in enumerate(info):

            table.cell(i,0).text = row[0]
            table.cell(i,1).text = row[1] or "-"



        # -------------------------
        # 3 Технические данные
        # -------------------------

        doc.add_heading(
            "3. Основные технические данные",
            level=2
        )


        table = doc.add_table(
            rows=8,
            cols=2
        )

        table.style = "Table Grid"


        technical = [

            (
                "Оборудование",
                data.get("equipment","")
            ),

            (
                "Мощность",
                data.get("power","")
            ),

            (
                "Номинальный ток",
                data.get("current","")
            ),

            (
                "Напряжение",
                data.get("voltage","")
            ),

            (
                "Степень защиты",
                data.get("ip","")
            ),

            (
                "Частота",
                data.get("frequency","")
            ),

            (
                "Серийный номер",
                data.get("serial_number","")
            ),

            (
                "Масса",
                data.get("weight","")
            )

        ]


        for i,row in enumerate(technical):

            table.cell(i,0).text = row[0]
            table.cell(i,1).text = row[1] or "-"



        # -------------------------
        # 4 Комплектность
        # -------------------------

        doc.add_heading(
            "4. Комплектность",
            level=2
        )


        items = [

            "Шкаф управления",

            "Паспорт оборудования",

            "Заводская документация",

            "Исполнительная документация ID-Agent"

        ]


        for item in items:

            doc.add_paragraph(
                "□ " + item
            )



        # -------------------------
        # 5 Ответственные лица
        # -------------------------

        doc.add_heading(
            "5. Ответственные лица",
            level=2
        )


        doc.add_paragraph(
            "Разработал: ____________________"
        )

        doc.add_paragraph(
            "Проверил: ______________________"
        )

        doc.add_paragraph(
            "Дата: __________________________"
        )



        doc.save(filename)


        return filename



executive_generator_v4 = ExecutiveDocumentGeneratorV4()
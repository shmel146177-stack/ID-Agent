import json
import os

from docx import Document
from docx.shared import Pt


class ProjectExecutiveGenerator:

    def create(self, project_name: str):

        project_path = os.path.join(
            "projects",
            project_name
        )

        analysis_path = os.path.join(
            project_path,
            "analysis"
        )

        executive_path = os.path.join(
            project_path,
            "executive_docs"
        )

        os.makedirs(
            executive_path,
            exist_ok=True
        )

        project_analysis_file = os.path.join(
            analysis_path,
            "project_analysis.json"
        )

        completeness_file = os.path.join(
            analysis_path,
            "document_completeness.json"
        )

        if not os.path.exists(project_analysis_file):
            raise FileNotFoundError(
                "project_analysis.json не найден"
            )

        with open(
            project_analysis_file,
            "r",
            encoding="utf-8"
        ) as file:
            project_analysis = json.load(file)

        completeness = {}

        if os.path.exists(completeness_file):

            with open(
                completeness_file,
                "r",
                encoding="utf-8"
            ) as file:
                completeness = json.load(file)

        documents = project_analysis.get(
            "documents",
            []
        )

        main_analysis = {}

        for document in documents:

            if document.get("status") == "Обработан":

                main_analysis = document.get(
                    "analysis",
                    {}
                )

                break

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

        # Титульная часть
        doc.add_heading(
            "ИСПОЛНИТЕЛЬНАЯ ДОКУМЕНТАЦИЯ",
            level=1
        )

        doc.add_paragraph(
            "Сформировано системой ID-Agent"
        )

        doc.add_paragraph(
            f"Проект: {project_name}"
        )

        doc.add_paragraph(
            "Объект: ______________________________"
        )

        doc.add_paragraph(
            "Адрес: _______________________________"
        )

        # 1. Общие сведения
        doc.add_heading(
            "1. Общие сведения",
            level=2
        )

        table = doc.add_table(
            rows=6,
            cols=2
        )

        table.style = "Table Grid"

        general_info = [
            (
                "Проект",
                project_name
            ),
            (
                "Оборудование",
                main_analysis.get("equipment", "")
            ),
            (
                "Изготовитель",
                main_analysis.get("manufacturer", "")
            ),
            (
                "Дата",
                main_analysis.get("date", "")
            ),
            (
                "Номер чертежа",
                main_analysis.get("drawing_number", "")
            ),
            (
                "Тип документа",
                main_analysis.get("document_type", "")
            )
        ]

        for i, row in enumerate(general_info):
            table.cell(i, 0).text = row[0]
            table.cell(i, 1).text = row[1] or "-"

        # 2. Технические характеристики
        doc.add_heading(
            "2. Основные технические характеристики",
            level=2
        )

        table = doc.add_table(
            rows=7,
            cols=2
        )

        table.style = "Table Grid"

        technical = [
            (
                "Мощность",
                main_analysis.get("power", "")
            ),
            (
                "Номинальный ток",
                main_analysis.get("current", "")
            ),
            (
                "Напряжение",
                main_analysis.get("voltage", "")
            ),
            (
                "Степень защиты",
                main_analysis.get("ip", "")
            ),
            (
                "Частота",
                main_analysis.get("frequency", "")
            ),
            (
                "Масса",
                main_analysis.get("weight", "")
            ),
            (
                "Серийный номер",
                main_analysis.get("serial_number", "")
            )
        ]

        for i, row in enumerate(technical):
            table.cell(i, 0).text = row[0]
            table.cell(i, 1).text = row[1] or "-"

        # 3. Реестр документов
        doc.add_heading(
            "3. Реестр документов проекта",
            level=2
        )

        registry_table = doc.add_table(
            rows=1,
            cols=4
        )

        registry_table.style = "Table Grid"

        headers = [
            "№",
            "Файл",
            "Тип",
            "Статус"
        ]

        for i, header in enumerate(headers):
            registry_table.cell(0, i).text = header

        for number, document in enumerate(
            documents,
            start=1
        ):

            row = registry_table.add_row().cells

            row[0].text = str(number)
            row[1].text = document.get(
                "filename",
                ""
            )
            row[2].text = document.get(
                "classification",
                "Не определён"
            )
            row[3].text = document.get(
                "status",
                ""
            )

        # 4. Комплектность
        doc.add_heading(
            "4. Контроль комплектности",
            level=2
        )

        doc.add_paragraph(
            f"Статус: {completeness.get('status', 'Нет данных')}"
        )

        doc.add_paragraph(
            f"Комплектность: "
            f"{completeness.get('completeness_percent', 0)}%"
        )

        doc.add_paragraph(
            f"Найдено: "
            f"{completeness.get('found_count', 0)} "
            f"из {completeness.get('required_count', 0)}"
        )

        completeness_table = doc.add_table(
            rows=1,
            cols=3
        )

        completeness_table.style = "Table Grid"

        completeness_table.cell(
            0,
            0
        ).text = "№"

        completeness_table.cell(
            0,
            1
        ).text = "Тип документа"

        completeness_table.cell(
            0,
            2
        ).text = "Наличие"

        for number, item in enumerate(
            completeness.get("documents", []),
            start=1
        ):

            row = completeness_table.add_row().cells

            row[0].text = str(number)

            row[1].text = item.get(
                "document_type",
                ""
            )

            row[2].text = item.get(
                "status",
                ""
            )

        # 5. Ответственные лица
        doc.add_heading(
            "5. Ответственные лица",
            level=2
        )

        doc.add_paragraph(
            "Разработал: __________________________"
        )

        doc.add_paragraph(
            "Проверил: ____________________________"
        )

        doc.add_paragraph(
            "Представитель заказчика: _____________"
        )

        doc.add_paragraph(
            "Дата: ________________________________"
        )

        filename = (
            f"Исполнительная_документация_{project_name}.docx"
        )

        file_path = os.path.join(
            executive_path,
            filename
        )

        doc.save(
            file_path
        )

        return file_path


project_executive_generator = ProjectExecutiveGenerator()